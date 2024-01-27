import argostranslate
from pdf2docx import Converter
import os
from docx import Document


def convert_pdf_to_docx(pdf_file, name):

    docx_file = name.split(".")[0] + ".docx"

    # Convert PDF to DOCX
    cv = Converter(pdf_file)
    cv.convert(docx_file, start=0, end=None)
    cv.close()

    return docx_file


def download_file(file):
    downloads_dir = os.path.join(os.path.expanduser('~'), 'Downloads')
    file_path = os.path.join(downloads_dir, file)
    os.rename(file, file_path)


def translate_text(docx, from_code, to_code):
   
#   from_code = "en"
#   to_code = "ru"
  
  docs = Document(docx)
  doc_text = list(set([ p.text for p in docs.paragraphs]))
  doc_text.remove('')
  
  for paragraph in docs.paragraphs:

    text = paragraph.text
      
    if text in doc_text:  
      if "•" in text:
        translation = argostranslate.translate.translate(text, from_code, to_code)
        paragraph.text = translation
      else:
    
        runs = iter(paragraph.runs)
        counter = 0
        
        for run in runs:
          if(counter == 0):
            run.text = argostranslate.translate.translate(text, from_code, to_code)
          else:
            run.text = ''
          counter+= 1
          #print(run.text)
          #print("Counter: ", counter)
          print("style: ", run.font.size)
          if "\t" in run.text:
            print("tab found!!!")
        
  for table in docs.tables:
    for row in table.rows:
        for cell in row.cells:
            for parag in cell.paragraphs:
              for run in parag.runs:
                run.text = argostranslate.translate.translate(run.text, from_code, to_code)
                
  docx_filename = docx.split(".")[0] + '_' + from_code +'_' + to_code + ".docx"#pdf_file.filename.split(".")[len(pdf_file.filename.split(".")) - 1]#'utput.docx'
  docs.save(docx_filename)

  return docx_filename